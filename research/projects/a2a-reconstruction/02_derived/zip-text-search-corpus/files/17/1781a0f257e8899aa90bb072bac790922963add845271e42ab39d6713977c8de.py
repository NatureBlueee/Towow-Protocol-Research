from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

src=Path('/mnt/data/unified_paper/output/Towow_A2A_Unified_Paper_raw.docx')
out=Path('/mnt/data/unified_paper/output/Towow_A2A_Unified_Paper.docx')
doc=Document(src)

BLUE=RGBColor(31,78,121)
GRAY=RGBColor(75,85,99)
LIGHT='F4F6F8'

# Page setup
for section in doc.sections:
    section.page_width=Cm(21.0)
    section.page_height=Cm(29.7)
    section.top_margin=Cm(2.2)
    section.bottom_margin=Cm(2.1)
    section.left_margin=Cm(2.35)
    section.right_margin=Cm(2.35)
    section.header_distance=Cm(0.9)
    section.footer_distance=Cm(0.9)

# Helpers

def set_run_font(run, name='Noto Serif CJK SC', size=None, bold=None, italic=None, color=None):
    run.font.name=name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.font.bold=bold
    if italic is not None: run.font.italic=italic
    if color is not None: run.font.color.rgb=color


def set_style_font(style_name, name, size, bold=None, color=None):
    style=doc.styles[style_name]
    style.font.name=name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    style.font.size=Pt(size)
    if bold is not None: style.font.bold=bold
    if color is not None: style.font.color.rgb=color
    return style

# Base styles
normal=set_style_font('Normal','Noto Serif CJK SC',10.5)
normal.paragraph_format.line_spacing=1.28
normal.paragraph_format.space_after=Pt(3)
normal.paragraph_format.first_line_indent=Pt(21)

for sn in ['Body Text','First Paragraph','Abstract','Bibliography']:
    st=set_style_font(sn,'Noto Serif CJK SC',10.5)
    st.paragraph_format.line_spacing=1.28
    st.paragraph_format.space_after=Pt(3)
    st.paragraph_format.first_line_indent=Pt(21)

compact=set_style_font('Compact','Noto Serif CJK SC',10.2)
compact.paragraph_format.line_spacing=1.16
compact.paragraph_format.space_after=Pt(1.5)
compact.paragraph_format.left_indent=Cm(0.7)
compact.paragraph_format.first_line_indent=Pt(0)

block=set_style_font('Block Text','Noto Serif CJK SC',10.2, color=GRAY)
block.paragraph_format.left_indent=Cm(0.8)
block.paragraph_format.right_indent=Cm(0.5)
block.paragraph_format.space_before=Pt(5)
block.paragraph_format.space_after=Pt(5)
block.paragraph_format.line_spacing=1.2
block.paragraph_format.first_line_indent=Pt(0)

# Title styles
title=set_style_font('Title','Noto Sans CJK SC',25,bold=True,color=BLUE)
title.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before=Pt(42)
title.paragraph_format.space_after=Pt(12)
sub=set_style_font('Subtitle','Noto Sans CJK SC',14,color=GRAY)
sub.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after=Pt(26)
for sn,sz in [('Author',11),('Date',10)]:
    st=set_style_font(sn,'Noto Sans CJK SC',sz,color=GRAY)
    st.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after=Pt(5)
abs_title=set_style_font('Abstract Title','Noto Sans CJK SC',12,bold=True,color=BLUE)
abs_title.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
abs=set_style_font('Abstract','Noto Serif CJK SC',9.5)
abs.paragraph_format.left_indent=Cm(0.8)
abs.paragraph_format.right_indent=Cm(0.8)
abs.paragraph_format.first_line_indent=Pt(19)
abs.paragraph_format.line_spacing=1.2

# Headings
for sn,sz,col,before,after in [
    ('Heading1',17,BLUE,18,8),
    ('Heading2',13.5,GRAY,13,5),
    ('Heading3',11.5,GRAY,9,3),
]:
    st=set_style_font(sn,'Noto Sans CJK SC',sz,bold=True,color=col)
    st.paragraph_format.space_before=Pt(before)
    st.paragraph_format.space_after=Pt(after)
    st.paragraph_format.keep_with_next=True
    st.paragraph_format.first_line_indent=Pt(0)

# Other styles
for sn in ['Caption','Image Caption','Table Caption']:
    st=set_style_font(sn,'Noto Sans CJK SC',9,color=GRAY)
    st.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_before=Pt(3)
    st.paragraph_format.space_after=Pt(8)
    st.paragraph_format.first_line_indent=Pt(0)
fig=set_style_font('Figure','Noto Serif CJK SC',10)
fig.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
code=set_style_font('Source Code','DejaVu Sans Mono',8.5)
code.paragraph_format.line_spacing=1.0
code.paragraph_format.space_before=Pt(4)
code.paragraph_format.space_after=Pt(4)

# Header and footer
def add_page_number(paragraph):
    run=paragraph.add_run()
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text=' PAGE '
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    set_run_font(run,'Liberation Sans',9,color=GRAY)

for section in doc.sections:
    hp=section.header.paragraphs[0]
    hp.text='从主权世界到共同现实  ·  通爻 Agent-to-Agent 统一论文'
    hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs: set_run_font(r,'Noto Sans CJK SC',8.5,color=GRAY)
    # subtle bottom border
    pPr=hp._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'4'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'D1D5DB')
    pBdr.append(bottom); pPr.append(pBdr)
    fp=section.footer.paragraphs[0]
    fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

# Paragraph-level cleanup and special formatting
for i,p in enumerate(doc.paragraphs):
    if p.style.name in ('Body Text','First Paragraph','Normal') and p.text.strip():
        p.paragraph_format.keep_together=False
        p.paragraph_format.widow_control=True
    if p.style.name=='Block Text':
        # shade quote paragraphs
        pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),LIGHT); pPr.append(shd)
    if p.style.name in ('Captioned Figure','Figure'):
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # Manual TOC and first main heading start on fresh pages
    if p.style.name=='Heading 1' and (p.text.strip()=='目录' or p.text.startswith('摘要说明')):
        p.paragraph_format.page_break_before=True
    if p.style.name=='Abstract Title' and not p.text.strip():
        p.text='摘要'
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: set_run_font(r,'Noto Sans CJK SC',12,bold=True,color=BLUE)
    # Appendices start on fresh pages
    if p.style.name=='Heading 1' and p.text.startswith('附录'):
        p.paragraph_format.page_break_before=True
    # References fresh page
    if p.style.name=='Heading 1' and p.text.startswith('参考文献'):
        p.paragraph_format.page_break_before=True

# Tables
for table in doc.tables:
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.style='Table'
    for r_idx,row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcPr=cell._tc.get_or_add_tcPr()
            margins=tcPr.first_child_found_in('w:tcMar')
            if margins is None:
                margins=OxmlElement('w:tcMar'); tcPr.append(margins)
            for side,val in [('top','70'),('left','90'),('bottom','70'),('right','90')]:
                node=margins.find(qn(f'w:{side}'))
                if node is None:
                    node=OxmlElement(f'w:{side}'); margins.append(node)
                node.set(qn('w:w'),val); node.set(qn('w:type'),'dxa')
            if r_idx==0:
                shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'DCE6F1'); tcPr.append(shd)
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0)
                p.paragraph_format.line_spacing=1.05
                p.paragraph_format.first_line_indent=Pt(0)
                for run in p.runs:
                    set_run_font(run,'Noto Serif CJK SC',8.3,bold=(r_idx==0))

# Set alt text from image filename where possible
for shape in doc.inline_shapes:
    try:
        docPr=shape._inline.docPr
        if not docPr.get('descr'):
            docPr.set('descr','通爻 Agent-to-Agent 统一论文示意图')
    except Exception:
        pass

# Core properties
doc.core_properties.title='从主权世界到共同现实：Agent-to-Agent 生成式协调的统一理论、协议与系统'
doc.core_properties.subject='通爻 Agent-to-Agent 统一论文'
doc.core_properties.author='通爻 Agent-to-Agent 研究项目'
doc.core_properties.keywords='Agent-to-Agent; 生成式协调; 主权边界; 可能性形成; Harness; WoWok'

doc.save(out)
print(out)
