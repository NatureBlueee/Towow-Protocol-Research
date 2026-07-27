from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

P=Path('/mnt/data/Towow_Unified_Paper_v1.0_formal/通爻_主权智能主体共同现实形成_正式论文_v1.0.docx')
doc=Document(P)
SERIF='Noto Serif CJK SC'
SANS='Noto Sans CJK SC'
MONO='Liberation Mono'


def set_run_font(run, name, size=None, bold=None, italic=None, color=None):
    run.font.name=name
    rpr=run._element.get_or_add_rPr()
    rfonts=rpr.rFonts
    if rfonts is None:
        rfonts=OxmlElement('w:rFonts'); rpr.insert(0,rfonts)
    for attr in ['ascii','hAnsi','eastAsia','cs']:
        rfonts.set(qn('w:'+attr),name)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic
    if color is not None: run.font.color.rgb=RGBColor(*color)


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    tc=cell._tc
    tcPr=tc.get_or_add_tcPr()
    tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None:
            node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')


def shade_paragraph(p, fill='F5F5F5'):
    pPr=p._p.get_or_add_pPr()
    shd=pPr.find(qn('w:shd'))
    if shd is None:
        shd=OxmlElement('w:shd'); pPr.append(shd)
    shd.set(qn('w:fill'),fill)
    shd.set(qn('w:val'),'clear')


def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr()
    tblHeader=trPr.find(qn('w:tblHeader'))
    if tblHeader is None:
        tblHeader=OxmlElement('w:tblHeader'); trPr.append(tblHeader)
    tblHeader.set(qn('w:val'),'true')


def prevent_row_split(row):
    trPr=row._tr.get_or_add_trPr()
    cant=trPr.find(qn('w:cantSplit'))
    if cant is None:
        cant=OxmlElement('w:cantSplit'); trPr.append(cant)


def add_page_field(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=paragraph.add_run()
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.extend([fldChar1,instr,fldChar2])
    set_run_font(run,SERIF,9,color=(90,90,90))

# Section layout, headers, footers
for sec in doc.sections:
    sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
    sec.top_margin=Cm(2.35); sec.bottom_margin=Cm(2.25)
    sec.left_margin=Cm(2.45); sec.right_margin=Cm(2.25)
    sec.header_distance=Cm(1.05); sec.footer_distance=Cm(1.05)
    sec.different_first_page_header_footer=True
    # regular header
    hp=sec.header.paragraphs[0]
    hp.text='主权智能主体的共同现实形成 · 正式论文 v1.0'
    hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs: set_run_font(r,SANS,8.5,color=(100,100,100))
    # first page header blank
    sec.first_page_header.paragraphs[0].text=''
    # regular footer page number
    fp=sec.footer.paragraphs[0]; fp.text=''; add_page_field(fp)
    sec.first_page_footer.paragraphs[0].text=''

# Title page separation: page break after Date paragraph (4th paragraph)
for p in doc.paragraphs[:10]:
    if p.style.name=='Date':
        p.add_run().add_break(WD_BREAK.PAGE)
        break

# Patch garbled TOC heading and request field update
ns={'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
sdts=doc._element.xpath('./w:body/w:sdt')
if sdts:
    tnodes=list(sdts[0].iter(qn('w:t')))
    if tnodes:
        tnodes[0].text='目录'
# updateFields in settings
settings=doc.settings._element
uf=settings.find(qn('w:updateFields'))
if uf is None:
    uf=OxmlElement('w:updateFields'); settings.append(uf)
uf.set(qn('w:val'),'true')

# Main paragraph formatting and chapter page breaks
front_h1={'摘要','Abstract','论文状态、主张强度与阅读约定'}
in_bibliography=False
for p in doc.paragraphs:
    txt=p.text.strip()
    sty=p.style.name
    if sty=='Heading 1':
        p.paragraph_format.page_break_before=True
        p.paragraph_format.keep_with_next=True
        p.paragraph_format.keep_together=True
        in_bibliography=(txt=='参考文献')
    elif sty.startswith('Heading'):
        p.paragraph_format.keep_with_next=True
        p.paragraph_format.keep_together=True
    if sty in ('Body Text','First Paragraph','Normal'):
        p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.widow_control=True
    if sty=='Compact':
        p.paragraph_format.widow_control=True
    if sty=='Source Code':
        p.paragraph_format.keep_together=True
        shade_paragraph(p,'F5F5F5')
        for r in p.runs: set_run_font(r,MONO,8.2)
    if p._p.xpath('.//m:oMathPara'):
        p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent=Cm(0)
        p.paragraph_format.space_before=Pt(4)
        p.paragraph_format.space_after=Pt(6)
    if in_bibliography and sty not in ('Heading 1','Heading 2','Heading 3') and txt:
        try:
            p.style=doc.styles['Bibliography']
        except KeyError:
            pass
        p.paragraph_format.left_indent=Cm(0.7)
        p.paragraph_format.first_line_indent=Cm(-0.7)
        p.paragraph_format.space_after=Pt(3)
        p.paragraph_format.line_spacing=1.15
        for r in p.runs: set_run_font(r,SERIF,9.2)

# Figure formatting and size cap
max_w=Cm(15.9)
for shape in doc.inline_shapes:
    if shape.width > max_w:
        ratio=max_w/shape.width
        shape.width=max_w
        shape.height=int(shape.height*ratio)
# Keep figure and caption together
for i,p in enumerate(doc.paragraphs):
    if p.style.name in ('Captioned Figure','Figure'):
        p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next=True
        p.paragraph_format.space_before=Pt(8)
        p.paragraph_format.space_after=Pt(2)
    elif p.style.name in ('Image Caption','Caption'):
        p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_together=True
        p.paragraph_format.space_after=Pt(8)
        p.paragraph_format.first_line_indent=Cm(0)

# Tables: compact but readable; repeat headers; prevent splitting rows
for tbl in doc.tables:
    tbl.autofit=True
    ncols=max((len(r.cells) for r in tbl.rows), default=1)
    fs=8.4 if ncols<=4 else (7.8 if ncols<=6 else 7.2)
    if tbl.rows:
        set_repeat_table_header(tbl.rows[0])
    for ri,row in enumerate(tbl.rows):
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcPr=cell._tc.get_or_add_tcPr()
            if ri==0:
                shd=tcPr.find(qn('w:shd'))
                if shd is None:
                    shd=OxmlElement('w:shd'); tcPr.append(shd)
                shd.set(qn('w:fill'),'E7EBEF'); shd.set(qn('w:val'),'clear')
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent=Cm(0)
                p.paragraph_format.space_before=Pt(0)
                p.paragraph_format.space_after=Pt(1.5)
                p.paragraph_format.line_spacing=1.05
                p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r,SERIF,fs,bold=(True if ri==0 else None))

# Make keywords compact and no indent
for p in doc.paragraphs:
    if p.text.strip().startswith(('关键词：','Keywords:')):
        p.paragraph_format.first_line_indent=Cm(0)
        p.paragraph_format.space_before=Pt(5)
        for r in p.runs: set_run_font(r,SERIF,9.5)

# metadata
cp=doc.core_properties
cp.title='主权智能主体的共同现实形成'
cp.subject='面向超级个体与一人公司的生成式协调理论、协议与运行系统'
cp.author='通爻研究计划'
cp.last_modified_by='通爻研究计划'
cp.comments='正式统一论文 v1.0；包含理论、形式模型、系统、实验与证据审计。'

doc.save(P)
print('saved',P)
